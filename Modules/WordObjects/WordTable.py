import docx
from enum import Enum
import copy
from lxml import etree
from copy import deepcopy
from docx.text.paragraph import Paragraph
import math
from Modules.WordObjects.InterfaceObj import WordObject, Gluing


class Table(WordObject):
    """ Параметры объекта: 
    title - массив строк с текстом, представляющих собой шапку таблицы
    record_sample - шаблон строки, полезно, если во всех строках есть объединение ячеек. Всегда массив массивов, даже если строка одна
    title и record_sample передаются в таком виде, в каком они должны быть в таблицы
    data = массив строк данных: массив массивов, где каждая строка соответствует record_sample
    vertical - маркер вертикальной таблицы (титул является столбцом)
    row_in_table - ограничение на количество строк в таблице"""
    
    def __init__(self, title = None, record_sample = None, is_vertical = None, row_in_table = None, footer = None, footer_rows = 0):
        
        if not title: self.title = None 
        # если передан одномерный массив, делаем из него двумерный
        elif self._get_matrix_size(title)[0] == 0: self.title = [title]
        else: self.title = title

        if not record_sample: self.record_sample = None
        elif self._get_matrix_size(record_sample)[0] == 0: self.record_sample = [record_sample]
        else: self.record_sample = record_sample
                
        self.vertical = is_vertical
        self.row_in_table = row_in_table

        if self.vertical: 
            if self.title: self.title = list(map(list, zip(*self.title)))
            if self.record_sample: self.record_sample = list(map(list, zip(*self.record_sample)))  


        self.title_style = {}
        self.record_style = {}

        self.data = None
        self.table_sample = None
        
        if not footer:
            self.footer_rows = footer_rows
            self.footer = [[]]
        else:
            if self._get_matrix_size(footer)[0] == 0: self.footer = [footer]
            else: self.footer = footer
            self.footer_rows = 0



    # выводит размер матицы в формате (кол-во строк, кол-во столбцов)
    def _get_matrix_size(self, arr):
        if all(not isinstance(i, list) for i in arr):  # Проверяем, одномерный ли список
            return (0,len(arr))  # Возвращаем количество элементов в одномерном списке
        else:
            return (len(arr), len(arr[0]) if arr else 0)  # Возвращаем количество столбцов (элементов в первом вложенном списке)


    """---------------копирование структуры таблицы--------------------------------------------------------------------------"""
    def scan_table(self,table): 
        # размер матрицы
        m = len(table.rows)
        n = len(table.columns)
        f = self.footer_rows
        
        table_map = [["" for j in range(n)] for i in range(m)]
        full_rows = [True for i in range(m)]
        full_cols = [True for i in range(n)]
        all_properties = {}


        #  сохраняем соотношение ширины столбцов
        columns_width = [0 for i in range(n)] #ширина каждого столбца
        columns_width_sum = 0 #ширина всей таблицы
        
        for i, cell in enumerate(table.columns):
            columns_width_sum += cell.width
            columns_width[i] = cell.width

        columns_width = [w/columns_width_sum for w in columns_width]
        self.record_style['COLUMNS_WIDTH'] = columns_width 

        
        # анализируем таблицу
        for i, row in enumerate(table.rows): 
            for j, cell in enumerate(row.cells): 
                
                properties = self._get_cell_properties(cell)
                # если табличка пустая, вычеркивам ее строку и столбец из списка заполненных
                if not cell.text: 
                    full_rows[i] = False
                    full_cols[j] = False

                # если в карту еще ничего не записано
                if not table_map[i][j]: 
                    #вносим в карту текст
                    table_map[i][j] = cell.text if cell.text else None
                    left_merged = 0
                    up_merged = 0 
                    
                    # если текущая ячейка объеденена с другими, сразу вносим все объединения в карту
                    #записываем объединения по горизонтали
                    if 'gridSpan' in properties.keys():
                        for k in range(j+1, j+int(properties['gridSpan'])):
                            table_map[i][k] = Gluing.LEFT
                            left_merged+=1
                        del properties['gridSpan']
                    
                    #записываем объединения по вертикали
                    if 'vMerge' in properties.keys():
                        cell_text = cell.text
                        k = i+1
                        #идем вниз по ячейкам, пока внизу есть объединение вверх 
                        while k < m and table.cell(k, j).text == cell_text and 'vMerge' in self._get_cell_properties(table.cell(k, j)):
                            table_map[k][j] = Gluing.UP
                            up_merged += 1
                            k += 1
                            
                        del properties['vMerge']
                    
                    # если оказалось, что ячейки объяденены в квадрат, дописываем диагонали
                    if left_merged and up_merged: 
                        for cell_i in range(i+1, i+up_merged + 1):
                            for cell_j in range(j+1, j+left_merged + 1): 
                                table_map[cell_i][cell_j] = Gluing.DIAGONAL
                    
                    #запоминаем прочие параметры ячеек 
                    for prop in properties.keys():
                        if prop in all_properties.keys():
                            all_properties[prop].append((i,j,properties[prop]))
                        else: 
                            all_properties[prop] = [(i,j,properties[prop])] 

        #не учитывам последние строчки
        if f: 
            full_rows = full_rows[:len(full_rows)-f]
            full_cols = full_cols[:len(full_cols)-f]
        
        # определяем направление таблицы
        if self.vertical == None:
            self.vertical = False if full_rows.count(True) >= full_cols.count(True) else True
        
        #переворачиваем, если таблица вертикальная
        if self.vertical: 
            table_map = list(map(list, zip(*table_map)))
            full_rows = full_cols
        
        # отделяем шапку таблицы
        k = 0
        while full_rows[k]: k+=1
        
        if not self.title:
            title = table_map[:k].copy()
            self.title = title
            table_map = table_map[k:]

        #отделяем футер таблицы
        if f: 
            if self.footer == None: 
                self.footer = table_map[len(table_map) - f :]
            table_map = table_map[:len(table_map) - f]
        
        #определяем шаблон записи таблицы
        if not self.record_sample:
            rows_in_record = None
            L = len(table_map)
            # ищим повторяющийся паттерн среди записей
            for k in range(1, L):
                if L%k == 0: 
                    rows_in_record = k
                    for i in range(k, L-k+1, k):
                        if table_map[i:i+k] != table_map[i-k:i]:
                            rows_in_record = None
                            break 
                if rows_in_record:
                    break
            
            if not rows_in_record: rows_in_record = L

            self.record_sample = table_map[0:rows_in_record]
        
            #записываем, сколько столбцов в таблице
            if self.vertical: self.row_in_table = L//rows_in_record

        title_size = self._get_matrix_size(self.title)
        record_size = self._get_matrix_size(self.record_sample)
 
        title_size = title_size[0]
        record_size = record_size[0]
          
        
        for key, value in all_properties.items():
            for element in value: 
                if self.vertical: 
                    element = (element[1], element[0], element[2])
                if element[0] < title_size: 
                    if key in self.title_style.keys():
                        self.title_style[key].append(element)
                    else: 
                        self.title_style[key] = [element]
                
                elif element[0]-title_size < record_size: 
                    element = (element[0]-title_size, element[1], element[2])
                    if key in self.record_style.keys():
                        self.record_style[key].append(element)
                    else: 
                        self.record_style[key] = [element]


        #находиим выравнивание текста в титульнике 
        if not 'ALIGNMENT' in self.title_style.keys():
            self.title_style['ALIGNMENT'] = []
            rows = table.rows if not self.vertical else table.columns
            find = False
            for i in range(0, title_size):
                row = rows[i]
                for j, cell in enumerate(row.cells): 
                    if cell.paragraphs: 
                        self.title_style['ALIGNMENT'].append(cell.paragraphs[0].alignment)
                    else: self.title_style['ALIGNMENT'].append(None)

        
        # находим выравнивание текста в записи
        if not 'ALIGNMENT' in self.record_style.keys():
            self.record_style['ALIGNMENT'] = []
            rows = table.rows if not self.vertical else table.columns
            find = False
            for i in range(title_size, title_size + record_size):
                row = rows[i]
                for cell in row.cells: 
                    if cell.paragraphs:
                        self.record_style['ALIGNMENT'].append(cell.paragraphs[0].alignment)
                    else:
                        self.record_style['ALIGNMENT'].append(None)

    #вытаскивает свойства ячейки из ее xml тега
    def _get_cell_properties(self, cell):
        properties = {}
        xml_cell = cell._element
        # Ищем тег w:tcPr (свойства ячейки)
        tc_pr = xml_cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if tc_pr is None:
            return properties
            
        for child in tc_pr.iter():
            tag = etree.QName(child).localname  # Убираем пространство имен из тега
            val = child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            
            if val is not None: 
                properties[tag] = val

        return properties
    
    
    """------------заполнение таблицы----------------------------------------------------------------------------------------"""
    def fill_table(self, table, data = None, delete_empty_rows=True): 
        if data == None: data = self.data
        if self._get_matrix_size(data)[0] >0:
            data = [cell for row in data for cell in row]
        data_index = 0
        rows = table.rows if not self.vertical else table.columns 
        cells_in_row = len(table.columns) if not self.vertical else len(table.rows)
        
        # если таблица не вертикальная, добавляем недостающие строки
        if not self.vertical:
            num_add_rows = math.ceil(len(data)/cells_in_row) - len(rows) + 1 + self.footer_rows

            row = rows[len(self.title)]
            self._add_n_rows(row, num_add_rows)

        merged_up = []
        number_of_empty_rows = 0
        # проходим по таблице и заполняем ее
        for i,row in enumerate(rows):
            for j,cell in enumerate(row.cells):
                if data_index >= len(data):
                    break
                else:
                    # добавляем в ячейку текст 
                    if data[data_index] not in Gluing: 
                        if data[data_index]== None: 
                            data_index +=1
                        elif cell.text == '':
                            # cell.text = str(data[data_index])
                            cell.paragraphs[0].text = str(data[data_index])
                            
                            data_index +=1
                    else:
                        # объединяем ячейки по горизонтали
                        if data[data_index] != Gluing.UP:
                            if self.vertical:
                                cell.merge(table.cell(j,i-1))
                            else:
                                cell.merge(table.cell(i,j-1))
                        
                        #объединение по вертикале произведем позже
                        else: 
                            merge = (i,j) if not self.vertical else (j,i)
                            merged_up.append(merge)
                        data_index +=1
            if data_index >= len(data):
                # сколько осталосб лишних строк
                number_of_empty_rows = len(rows) - i-1 - self.footer_rows
                break

        #удаляем лишние столбцы/строки если надо
        if delete_empty_rows and number_of_empty_rows:
            if self.vertical:
                self._delete_last_n_columns(table, number_of_empty_rows)
            else:
                self._delete_last_n_rows(table,number_of_empty_rows)
        
        #объединяем ячейки вверх
        for i,j in merged_up: 
            if i>0:
                table.cell(i,j).merge(table.cell(i-1,j))

    # копирует указанную строку n раз 
    def _add_n_rows(self, row, n):
        tr = row._tr
        for i in range(n): 
            new_tr = deepcopy(tr)
            tr.addnext(new_tr)   

    # удаляет последние n столбцов таблицы
    def _delete_last_n_columns(self, table, n):
        if n <= 0: 
            return None
        if len(table.columns) <= n: 
            return None

        m = len(table.columns) - n
        for i in range(n):
            for row in table.rows: 

                cell = row.cells[-1]
                prop = self._get_cell_properties(cell)    
                span = prop['gridSpan'] if 'gridSpan' in prop.keys() else None          

                if span: 
                    xml_cell = cell._element
                    tc_pr = xml_cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    grid_span = tc_pr.find('.//w:gridSpan', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if int(span) == 2:
                        tc_pr.remove(grid_span)
                    elif int(span) > 2: 
                        grid_span.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(span)-1))

                else: 
                    parent = cell._tc.getparent()
                    parent.remove(cell._tc)      
    
    # удаляет последние n строк таблицы
    def _delete_last_n_rows(self, table, n): 

        for _ in range(n): 
            row = table.rows[-1]
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

    """------------создание таблицы------------------------------------------------------------------------------------------"""              
    def _create_table(self, doc, data):
        rows = len(data)
        cols = len(data[0])
        if not self.vertical:
            table = doc.add_table(rows, cols, 'Table Grid')
            N = rows
        else:
            table = doc.add_table(cols, rows, 'Table Grid')
            N = cols
        # заполняем таблицу
        self.fill_table(table, data, True)
        
        # накладываем на таблицу настройки стилей
        # применяем стили для шапки таблицы
        if self.title_style and self.title:
            style = self.title_style.copy()
            
            if 'ALIGNMENT' in style.keys():

                alignments = style['ALIGNMENT']
                rows = table.rows if not self.vertical else table.columns
                k = 0
                for i in range(0, len(self.title)):
                    row = rows[i]
                    for j, cell in enumerate(row.cells):
                        if cell.text: 
                            cell.paragraphs[0].alignment = alignments[k]
                        k+=1               
                del style['ALIGNMENT']
            
            for key in style.keys():
                elements = style[key]
                for el in elements: 
                    i, j, val = el
                    if self.vertical: i,j = j,i
                    self._add_cell_property(table.cell(i,j), key, val)
        
        # приименяем стили ко всей остальной таблице
        if self.record_style:
            style = self.record_style.copy()
            # задаем ширину столбцов
            if 'COLUMNS_WIDTH' in style.keys():
                table.autofit = False
                width = 0
                
                columns_width = style['COLUMNS_WIDTH']
                for column in table.columns:
                    width += column.width
                for i,column in enumerate(table.columns): 
                    col_width = width*columns_width[i]
                    for cell in column.cells:
                        cell.width = col_width
                del style['COLUMNS_WIDTH']

            if 'ALIGNMENT' in style.keys():
                align = style['ALIGNMENT']
                rows = table.rows if not self.vertical else table.columns
                # n - количество строк размера record_sample в таблице
                n = self.row_in_table if self.row_in_table else (len(rows)-len(self.title))//len(self.record_sample)
                alignments = align*n #выравнивания каждой ячейки строк
                k = 0
                # выставляем выравнивание каждой ячейке
                for i in range(len(self.title), len(rows)):
                    row = rows[i]
                    for j, cell in enumerate(row.cells):
                        if cell.text: 
                            cell.paragraphs[0].alignment = alignments[k]
                        k+=1
                del style['ALIGNMENT']

                rows = table.rows if not self.vertical else table.columns
                for key in style.keys():
                    elements = style[key]
                    for el in elements: 
                        i, j, val = el
                        for i in range(len(self.title), len(rows), len(self.record_sample)): 

                            if self.vertical:
                                self._add_cell_property(table.cell(j,i), key, val)
                            else:
                                self._add_cell_property(table.cell(i,j), key, val)
        return table

    def _add_cell_property(self, cell, property, val):
        # Получаем XML элемент ячейки
        xml_cell = cell._element
        
        # Ищем элемент w:tcPr (свойства ячейки)
        tc_pr = xml_cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        if tc_pr is None:
            # Если элементов w:tcPr нет, создаем его
            tc_pr = etree.SubElement(xml_cell, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')

        # Проверяем, существует ли уже элемент w:textDirection
        tag = './/w:' + property
        prop = tc_pr.find(tag, namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        if prop is None:
            # Если элемента нет, создаем его
            prop = etree.SubElement(tc_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + property)
        
        # Устанавливаем значение атрибута w:val
        prop.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', val)


        # return table

    """------------полное копирование таблицы--------------------------------------------------------------------------------"""
    #полностью скопировать шаблон таблицы
    def copy_table(self, table, delete_after_copy = False): 
        tbl = table._tbl
        self.table_sample = deepcopy(tbl)
        self.scan_table(table)

        if delete_after_copy: 
            tbl.getparent().remove(tbl)

    #вставить и заполнить шаблон
    def _add_copy_table(self, doc, data): 
        p = doc.add_paragraph()
        p = p._p
        new_tbl = deepcopy(self.table_sample)
        p.addnext(new_tbl)
        p.getparent().remove(p)  # Удаление из родительского элемента

        table = doc.tables[-1]
        self.fill_table(table,data,True)
        return table
    
    '''-------------форматный ввод данных-----------------------------------------------------------------------------------'''
    def write_data(self, data): #внесение строк данных в таблицу 
        if not self.record_sample and not self.title: 
            self.data = data
            return data
        
        if not self.record_sample: 
            self.record_sample = [[None for i in range(len(self.title[0]))]]

        record_sample = self.record_sample
        
        # количество None в шаблоне (т.е. кол-во значений для заполнения)
        nums = sum(1 for row in record_sample for item in row if item is None)

        data_size = self._get_matrix_size(data)
        if data_size[0] == 0:
            if data_size[1] > nums:
                self.data = data
                return data
            else:
                data = [data]

        rows, cols = self._get_matrix_size(record_sample)
        rows = 1 if not rows else rows
        table_rows = []
        for data_row in data: 
            table_row = []
            # если в данных нет форматирования и количество переменных правильное, преобразуем строку данных по шаблону
            if len(data_row) <= nums and not any(item in data_row for item in Gluing):
                if len(data_row) < nums: 
                    data_row += ['' for i in range(nums-len(data_row))]
                table_row = copy.deepcopy(record_sample)
                k = 0
                for i in range(rows):
                    table_row_i = table_row[i]
                    for j in range(cols):
                        #вместо None в шаблоне вписываем данные  
                        if table_row_i[j] == None: 
                            table_row_i[j] = data_row[k]
                            k+=1
            # если строка неформатная, но вписывается в таблицу, записываем ее прямо так
            elif len(data_row)%cols == 0:
                table_row = [data_row.copy()] 
            
            # заносим строку в матрицу данных
            if table_row: table_rows+= table_row 
        self.data = table_rows
    
    """-------------печеть таблицы в документ-------------------------------------------------------------------------------"""
    def add_to_file(self, doc, add_after = None):
        Data = self.data.copy()

        if not self.row_in_table:
            tables = [Data]
        else: #если есть ограничение на кол-во строк в таблице, разбиваем ее
            k = 0
            num_rows = 0
            tables = []
            while len(Data)>0:                 
                num_rows += 1 

                if self.record_sample: 
                    k += len(self.record_sample) - 1
            
                else:
                    # если есть объединение считаем все одной строкой
                    repeat = True 
                    while repeat: 
                        repeat = False
                        if k+1<len(Data):
                            if not self.vertical and (Gluing.UP in Data[k+1]):
                                repeat = True
                                k+=1
                            elif self.vertical and (Gluing.LEFT in Data[k+1]):
                                repeat = True
                                k+=1
                # отщипляем row_in_table строк в отельную таблицу
                if num_rows == self.row_in_table:
                    
                    if Data[:k+1]:
                        tables.append(Data[:k+1].copy())
                    Data = Data[k+1:]
                    num_rows = 0
                    k= 0
                else:
                    k+=1                
            if Data: 
                tables.append(Data.copy())

        # создаем таблицу/ы
        for table in tables:
            if self.table_sample is None: #создаем таблицу
                new_table = self._create_table(doc, self.title + table + self.footer) 
            else: #или копируем шаблон 
                new_table = self._add_copy_table(doc,table)
            
            if add_after: #перетаскиваем таблицу, если указано куда
                add_after = self._add_table_after(add_after, new_table)
            else:
                add_after = doc.add_paragraph()
        return add_after
    
    # перемещает таблицу на место после add_after
    def _add_table_after(self, add_after, table):

        p_before = add_after._p
        tbl = table._tbl
        new_tbl = deepcopy(tbl)
        p_before.addnext(new_tbl)
        t_element = table._element  # Доступ к XML-элементу
        t_element.getparent().remove(t_element)  # Удаление из родительского элемента
        
        new_paragraph = add_after._parent.add_paragraph()  # Создаем новый параграф
        new_p = new_paragraph._p  # Доступ к внутреннему элементу <w:p>
    
        # Вставляем пустой параграф после таблицы
        new_tbl.addnext(new_p)            
        # return Paragraph(new_p, doc)
        return Paragraph(new_p, new_p.getparent())
    

if __name__ == "__main__":

    A = Table()
    # A.vertical = False
    A.footer_rows = 1
    path = "Shablon_new.docx"
    doc = docx.Document(path)
    table = doc.tables[2]
    A.copy_table(table)
    # print(A.title)
    # print(A.record_sample)
    # print(A.footer)
    Data = [[1,2,3,4,5,6],[1,2,3,4,5,6],[1,2,3,4,5,6],[1,2,3,4,5,6]]
    Data = [["(Импортные) Конденсаторы керамические постоянной емкости общего применения (CK, CKR)", Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT],
            ["C1", "C2012X7R2A104 (SMD0805 X7R 0,1 мкФ 100 В)", "0.99", "3.4787", "7", "58.8132"],
            ["C2", "C1608X7R1H104 (SMD0603 X7R 0,1 мкФ 50 В)", "0.99", "3.4787", "7", "58.8132"],
            ["C3", "C1608X7R1H104 (SMD0603 X7R 0,1 мкФ 50 В)", "0.99", "3.4787", "7", "58.8132"]]
    A.write_data(Data)
    # print(A.data)
    A.add_to_file(doc)

    # data2 = [["(Импортные) Логические схемы, программируемые логические матрицы, микросхемы памяти, микропроцессоры", 
    #             Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT],
    #         ["D1", "LM26CIM5-SHA", "1000", "20.6423", "1", "4", "13.4628"],
    #         ["D4", "LM26CIM5-VHA", "1000", "20.6423", "1", "4", "13.4628"],
    #         ["(Импортные) Полупроводниковые аналоговые", Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT, Gluing.LEFT],
    #         ["D2", "FODM8801C", "1000", "7.4865", "1", "4", "5.0531"],
    #         ["D3", "FODM8801C", "1000", "7.4865", "1", "4", "5.0531"]]

    # B = Table()
    # B.footer_rows = 1
    # table = doc.tables[4]
    # B.copy_table(table)
    # B.write_data(data2)
    # B.add_to_file(doc)
    # p = doc.paragraphs[1]
    # i = 0
    # for i,p in enumerate(doc.paragraphs):
    #     if p.text == "4. ВЫВОДЫ":
    #         break
    # print(p.text, i)



    doc.save("тесты_шаблон.docx")

    # doc = docx.Document(path)
    # A = Table()
    # A.vertical = True
    # path = "тесты.docx"
    # doc = docx.Document(path)
    # sample_doc = docx.Document("Гост.docx")
    # table = sample_doc.tables[-14]

        

    # A.scan_table(table)
    # # A.copy_table(table)
    # print(A.vertical)
    # print(A.title)
    # print(A.record_sample)
    # # A._add_copy_table(doc,['']) 
    # data = [['name 1'],['name 2'],['name 3']]

    # A.write_data(data)

    # A.add_to_file(doc)


    # doc.save("тесты2.docx")