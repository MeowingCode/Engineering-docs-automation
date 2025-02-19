import docx
from enum import Enum
import copy
import re

from copy import deepcopy
from docx.text.paragraph import Paragraph
from Modules.WordObjects.InterfaceObj import WordObject 

class TextBlock(WordObject):
    #  Параметры:
    # mark - слово, по которому модуль определяет начало и конец блока 
    # words_to_replace - словарь слов на замену 

    # self.textblock - массив параграфов 
    def __init__(self, doc, start, end, words_to_replace = None):
        
        if words_to_replace:
            self.replace = words_to_replace.copy()
        else: self.replace = None
        
        # ищем блок
        self.textblock = []
        paragraphs = doc.paragraphs
        # start = -1
        # end = -1
        # for i, pr in enumerate(paragraphs):
        #     if pr.text.strip() == ("__" + mark):
        #         start = i
        #     if pr.text.strip() == (mark + "__"):
        #         end = i
        #         break
        # if end<=start or start<0:
        #     return None 

        #сохраняем блок
        for i in range(start, end):
            self.textblock.append(deepcopy(paragraphs[i]._p))

        # #удаляем этот блок из файла
        for i in range(end, start-1,-1):
            p_element = paragraphs[i]._element  # Доступ к XML-элементу
            p_element.getparent().remove(p_element)  # Удаление из родительского элемента    

    def change_words_to_replace(self, words_to_replace):
        self.replace = words_to_replace.copy()

    #ищет все в параграфе, что надо заменить
    def get_marks_from_paragraph(self, paragraph):
        if not self.replace:
            return None 
        pattern = r'__(.*?)__'
        markers = re.findall(pattern, paragraph.text)
        
        marks = [mark for mark in markers if mark in self.replace.keys()]
        if marks:
            return marks
        else:
            return None
    
    # подставляет значения в место для замены в конкретном параграфе
    def replace_marks(self, paragraph):
        marks = self.get_marks_from_paragraph(paragraph)
        if not marks:
            return None   
        for mark in marks:
            for run in paragraph.runs:
                # print(run.text, end='!')
                run.text = run.text.replace("__" + mark + "__", self.replace[mark])
        # print('\n','_'*50)

    #печатает блок в указанное место (doc/таблица)
    def add_to_file(self, place_to_add, add_after = None):
        if len(self.textblock) > 1:
            last_paragraph = self.textblock[-1]
        elif len(self.textblock) == 1:
            last_paragraph = self.textblock[0]
        else:
            return None
        
        if not add_after:
            if len(place_to_add.paragraphs) > 1:
                add_after = place_to_add.paragraphs[-1]
            elif len(doc.paragraphs) == 1:
                add_after = place_to_add.paragraphs[0]
        
        add_after = add_after._p
        last_p = deepcopy(last_paragraph)
        add_after.addnext(last_p)

        last_paragraph = Paragraph(last_p, place_to_add)

        self.replace_marks(last_paragraph)
        if len(self.textblock) == 1:
            return last_paragraph
        
        for i in range(0, len(self.textblock)-1):
            new_paragraph = last_paragraph.insert_paragraph_before()
            
            # Очищаем содержимое нового параграфа (по умолчанию он пуст, но на всякий случай)
            new_paragraph.clear()

            # Добавляем содержимое из self.textblock[i]
            for child in deepcopy(self.textblock[i]):
                new_paragraph._p.append(child)
            
            self.replace_marks(new_paragraph)

        return last_paragraph
    
if __name__ == "__main__":
    PATH = "Базовые_тесты.docx"
    doc = docx.Document(PATH)   
    replace = {'ЗАМЕНА1': 'МЯУ','ЗАМЕНА2': 'мяу'} 
    A = TextBlock(doc,"ТЕСТ", replace)
    replace2 = {'ЗАМЕНА1': 'MEOW','ЗАМЕНА2': 'meow'} 
    
    A.add_to_file(doc)

    A.change_words_to_replace(replace2)
    A.add_to_file(doc)
    doc.save("тесты.docx")

"""заметки:
1) вставка в пустой файл работает не супер хорошо, добавляется после пустого параграфа. Не критично и случай редкий, но можно
исправить"""