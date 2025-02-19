import re
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Pt
from copy import deepcopy
from docx.text.paragraph import Paragraph
from Modules.WordObjects import Counter as C
from Modules.WordObjects.InterfaceObj import WordObject 


# класс для обычного текста 

class TextRow(WordObject):
        #параметры:
        # text - массив команд (b, i, u и т.д) и текста
        # headline_lv - уровень заголовка, если нужен заголовок, а не обычный текст
        # alignment - выравнивание текста 
        # font_size - размер текста
        # font_name - шрифт 

    def __init__(self, text, headline_level = 0, alignment = None, font_size = None, font_name = None):

        self._tags = {"b": False, "i": False, "u": False, "s": False, "sub": False, "sup": False}
        # b - bold - жирный шрифт 
        # i - italic - курсив 
        # u - подчеркивание
        # s - зачеркнутый
        # sub - подстрочный текст
        # sup - надстрочный текст
        
        self._one_tags = {}
        self._text = self._split_by_tags(text)  # преобразует в массив команд и текста
        
        # устанавливает уровень заголовка

        if headline_level and str(headline_level).isdigit() and int(headline_level) > 0:
            self._headline_lv = int(headline_level)
        else:
            self._headline_lv = 0

        # устанавливаем выравнивание
        if alignment == 'left': self._alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'center': self._alignment = WD_ALIGN_PARAGRAPH.CENTER 
        elif alignment == 'right': self._alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        elif alignment == 'justify': self._alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
        else: self._alignment = None

        #устанавливаем размер текста
        if font_size and str(font_size).isdigit() and int(font_size) > 0:
            self._font_size = Pt(int(font_size))
        else:
            self._font_size = None

        #устанавливаем название шрифта
        if font_name:
            self._font_name = str(font_name)
        else: self._font_name = None
        
        self.counter = C.Counter()

    # превращает текст с блоками в массив из текста и команд
    # например: "<b>Te<d>xt</b>" в массив ['b', 'Te<d>xt', '/b']
    def _split_by_tags(self, text):

        one_tag_pattern = r"(<[^<>]*?/>)"
        
        text = re.split(one_tag_pattern, text)
        text = [el for el in text if el]
        self._one_tags = {}
        
        for j, el in enumerate(text):
            # находим элемент с одиночным тегом
            if el[0] == '<' and el[-2:] == '/>':
                tag = el[1:-2].split(" ")
                tag_name = tag[0] + '_' + str(len(self._one_tags))
                params = {'tag': tag[0]}
                # записываем параметры
                for par in tag[1:]:
                    par = par.replace('"', '').replace("'", "")
                    parametr = par.split("=")
                    if len(parametr)>1: params[parametr[0]] = parametr[1] 

                # добавляем запись об этом одиночном теге
                self._one_tags[tag_name] = params.copy()
                # заменяем элемент с тегом на внутреннее название
                text[j] = tag_name


        # result = re.split(one_tag_pattern, text)
        tags = self._tags.keys()
        # Создаем шаблон для тегов, который ищет открывающие и закрывающие теги
        tag_pattern = r'(</?(?:' + '|'.join(tags) + r')>)'

        for i in range(len(text)-1, -1,-1):
            el = text[i]
            new_el = [el for el in re.split(tag_pattern, el) if el]
            text = text[:i] + new_el + text[i+1:]

        # # Убираем пустые строки и None
        # result = [elem for elem in result if elem and elem]

        # Преобразуем теги вида "<b>" в "b", и "</b>" в "/b"
        for i in range(len(text)):
            if text[i].startswith('<') and text[i].endswith('>'):
                text[i] = text[i][1:-1]  # убираем "<" и ">"
        return text
    
    #добавляет текст в документ 
    def add_to_file(self, place_to_add, add_after = None):
        
        # создаем параграф для текста
        if not self._headline_lv:
            # создаем параграф
            p = place_to_add.add_paragraph()
        else:
            # или создаем заголовок
            p = place_to_add.add_heading("", self._headline_lv)
        

        # применение команд к тексту по каждому рану
        for element in self._text:
            
            # если element - текст для вставки, то применяем к нему настройки run-а
            if element.lstrip('/') not in self._tags:
                if element in self._one_tags.keys():                
                    params = self._one_tags[element]

                    if params['tag'] == 'number':
                        name = params['name'] if 'name' in params.keys() else None
                        format = params['format'] if 'format' in params.keys() else None
                        num = self.counter.get_number(p,name,format)
                        run = p.add_run(num)
                else:
                    run = p.add_run(element)
                run.bold = self._tags["b"]
                run.italic = self._tags["i"]
                run.font.strike = self._tags["s"]
                run.underline = self._tags["u"]
                run.font.subscript = self._tags["sub"]
                run.font.superscript = self._tags["sup"]
            else:
                # если element - команда открытия тега, то обновляем флаги
                if element in self._tags:
                    self._tags[element] = True
                    # если element - команда закрытия тега, то наоборот убираем флаг
                else: self._tags[element.lstrip('/')] = False
        
        # если есть, добавляем выравнивание
        if self._alignment:
            p.alignment = self._alignment
        
        # изменяем высоту текста
        if self._font_size:
            try:
                for run in p.runs:
                    run.font.size = self._font_size
            except:
                pass
        
        # изменяем шрифт
        if self._font_name:
            try:
                for run in p.runs:
                    run.font.name = self._font_name
            except:
                pass
        
        # возможно заменю на это
        # style = p.style
        # if style:
        #     style.font.name = self._font_name
        #     style.font.size = self._font_size

        if add_after:
            p_before = add_after._p
            new_p = deepcopy(p._p)
            p_before.addnext(new_p)
            p_element = p._element  # Доступ к XML-элементу
            p_element.getparent().remove(p_element)  # Удаление из родительского элемента
            return Paragraph(new_p, place_to_add)
        return p


if __name__ == "__main__":
    str1 = "Я считаю <b> до <number name='Num0' format='num'/> не могу до <number name='Num0' format='num'/></b>"
    str2 = "<b> hr <number name='header' format='format'/>Узел 'НКГД.306559.528 Модуль светодиодный ССД600'</b>"
    A = TextRow(str2, "-104s", None, 12, 'Calibri')
    print(A._split_by_tags(str2))


