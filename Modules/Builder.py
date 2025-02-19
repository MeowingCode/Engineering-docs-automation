import docx
from copy import deepcopy

import Modules.WordObjects as WordObjects
from Modules.WordObjects import Replace

#создает заполняет документ 
#управляет всеми моими word_объектами 

class DocBuilder:
    def __init__(self, doc_path, table_sample_path):
        self.doc_path = doc_path 
        self.table_sample_path = table_sample_path
        
        self.add_after = None 
        self.table_samples = []
        self.textblock_samples = []

        self.DELIMITER = "__"
        self.FUNC_WORDS = ("TEXT", "TABLE")
        self.default_table_values = (None, None, False, None, None, 1)

        if not doc_path.is_file():
            self.doc = docx.Document()
        else:
            self.doc = docx.Document(self.doc_path)


    def _element_index(self, doc): 
        element_indexes = {}
        paragraphs = []
        tables = []
        for i, el in enumerate(doc.element.body):
            if el.tag.endswith('p'):
                paragraphs.append(i)
            elif el.tag.endswith('tbl'):
                tables.append(i)
        
        element_indexes["paragraphs"] = paragraphs
        element_indexes["tables"] = tables
        return (paragraphs,tables)

    def read_samples(self):
        DELIMITER = self.DELIMITER
        FUNC_WORDS = self.FUNC_WORDS
        default_table_values = self.default_table_values
        doc = self.doc
        # шаблоны из шаблона документа
        table_samples = {}
        text_block_samples = {}

        # записываем индексы таблиц и параграфов относительно друг друга
        p_index, tab_index = self._element_index(doc)

        end_text_block_index = -1    

        for p_i in range(len(doc.paragraphs)-1,-1,-1):
            p = doc.paragraphs[p_i]
            # если начинается с служебного символа 
            if p.text.startswith(DELIMITER): 
                words = [word for word in p.text.split(DELIMITER) if word]

                # если блок TEXT
                if words[0] == FUNC_WORDS[0]: 
                    if len(words) == 1:
                        end_text_block_index = p_i
                    elif end_text_block_index > p_i: 
                        #создаем объект шаблона
                        textblock = WordObjects.TextBlock(doc, p_i + 1, end_text_block_index)
                        text_block_samples[words[1]] = deepcopy(textblock)
                
                # если блок TABLE
                if words[0] == FUNC_WORDS[1]:
                        table = None
                        index = p_index[p_i]
                        
                        #индекс таблицы после тега 
                        j = next((j for j in range(len(tab_index)) if tab_index[j] > index), None)

                        if j is not None:
                            #создаем объект шаблона
                            Tab = WordObjects.Table(*default_table_values)
                            table_sample = doc.tables[j]
                            Tab.copy_table(table_sample, True)
                            table_samples[words[1]] = deepcopy(Tab)

                #блоки удаляются внутри объектов 
                #удаляем тег блока
                p._p.getparent().remove(p._p)
        self.table_samples = table_samples
        self.textblock_samples = text_block_samples       
        self.doc.save(self.doc_path)

    def replace_words(self, words_to_replace): 
        self.doc.save(self.doc_path)
        Replace.replace_words(self.doc_path,words_to_replace)
        self.doc = docx.Document(self.doc_path)

    def set_add_after(self, add_after_text): 
        if add_after_text is None: 
            self.add_after = None
        else: 
            self.add_after = next((p for p in self.doc.paragraphs if add_after_text in p.text), None)

    def add_paragraph(self, text, headline_level, alignment, font_size, font_name):
        headline_level = 0 if headline_level == None else headline_level
        row = WordObjects.TextRow(text,headline_level, alignment, font_size, font_name)
        if self.add_after is None:
            row.add_to_file(self.doc)
        else:
            self.add_after = row.add_to_file(self.doc, self.add_after)
        self.doc.save(self.doc_path)

    def add_textblock(self, sample_name:str, replace_dict:dict):
    
        if sample_name in self.textblock_samples.keys():
            block = self.textblock_samples[sample_name]
            block.change_words_to_replace(replace_dict.copy())
            if self.add_after is None:
                block.add_to_file(self.doc)
            else:
                self.add_after = block.add_to_file(self.doc, self.add_after)
        self.doc.save(self.doc_path)

    def add_table(self, mode, mark, vertical, table_title, record_sample, footer, table_data):
        
        if mode=="delete" and mark.isdigit() and int(mark)>0:
            num = int(mark)-1   
            if len(self.doc.tables)<=num:
                return
            table = self.doc.tables[num]
            tbl = table._tbl 
            parent = tbl.getparent()
            tbl.getparent().remove(tbl)
            

            self.doc.save(self.doc_path) 
            return 
                
        elif mode=="fill" and mark.isdigit() and int(mark)>0: 
            num = int(mark)-1   
            doc_tables = self.doc.tables
            if len(self.doc.tables)<=num:
                return
            doc_table = self.doc.tables[num]
            table = WordObjects.Table()
            table.scan_table(doc_table)
            table.fill_table(doc_table, table_data)
            self.doc.save(self.doc_path) 
            return 
        
        elif mode=="sample" and mark in self.table_samples.keys():
            table = self.table_samples[mark]
        # если не указан мод, создаем таблицу
        else: 
            #  title = None, record_sample = None, is_vertical = None, row_in_table = None, footer = None, footer_rows = 0):
            table = WordObjects.Table(table_title, record_sample, vertical, None, footer, None)
        table.write_data(table_data)     
        
        # запись таблицы в файл
        if self.add_after is None:
            table.add_to_file(self.doc)
        else:
            self.add_after = table.add_to_file(self.doc, self.add_after)
        self.doc.save(self.doc_path)        
        pass